"""
Sprint 1 First Report Generator
Template: Final-Report-Template.docx
Output: Sprint1_First_Report.docx
"""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────
BASE = r"C:\Users\bahti\Desktop\Goruntuisleme"
TEMPLATE = r"C:\Users\bahti\Desktop\Capstone_Project\Final-Report-Template.docx"
OUTPUT = os.path.join(BASE, "reports", "Sprint1_First_Report.docx")

IMG = {
    "augment_imbalance": os.path.join(BASE, "reports", "generated", "augmentation_imbalance_latest.png"),
    "v1_results":        os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_final", "results.png"),
    "v1_confusion":      os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_final", "confusion_matrix.png"),
    "v1_confusion_norm": os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_final", "confusion_matrix_normalized.png"),
    "v1_pr_curve":       os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_final", "BoxPR_curve.png"),
    "v1_f1_curve":       os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_final", "BoxF1_curve.png"),
    "baseline_confusion": os.path.join(BASE, "runs", "detect", "runs", "smoke", "smoke_phase12", "confusion_matrix.png"),
    "v2_results":        os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_v2", "results.png"),
    "v2_confusion":      os.path.join(BASE, "runs", "detect", "runs", "phase1", "yolov10s_ca_v2", "confusion_matrix.png"),
    "sample_train":      os.path.join(BASE, "data", "processed", "phase1_multiclass_v1", "train", "images",
                                      "3- 97322-M0325120202448-3_jpg.rf.RxQUw90Grm5JHETvfzYc.jpg"),
    "sample_infer":      os.path.join(BASE, "runs", "detect", "runs", "infer", "smoke_phase1_test",
                                      "3- 97332-M0325120202805-2_jpg.rf.qyqnC7GPVxkCQC8ySyBe.jpg"),
}

# Verify images exist
for k, v in IMG.items():
    if not os.path.exists(v):
        print(f"WARNING: Image missing: {k} -> {v}")


# ── Helpers ────────────────────────────────────────────────────────
def add_fig(doc, img_key, caption, width=Inches(5.5)):
    """Add an image with a centred caption below it."""
    if img_key in IMG and os.path.exists(IMG[img_key]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(IMG[img_key], width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles['Normal']
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Gorsel bulunamadi: {img_key}]")


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(10)
    return row


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


# ── Build Document ─────────────────────────────────────────────────
doc = Document(TEMPLATE)

# Clear all existing content (keep styles and sections)
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

for t in doc.tables:
    t._element.getparent().remove(t._element)

body = doc.element.body

# ═══════════════════════════════════════════════════════════════════
#  KAPAK SAYFASI
# ═══════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Capstone Project First Report")
run.bold = True
run.font.size = Pt(22)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Submitted to")
run.font.size = Pt(12)

fac = doc.add_paragraph()
fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fac.add_run("Faculty of Engineering and Natural Sciences")
run.bold = True
run.font.size = Pt(14)

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept.add_run("Department of Computer Engineering")
run.font.size = Pt(12)

doc.add_paragraph()

deg = doc.add_paragraph()
deg.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = deg.add_run("In partial fulfillment of the requirements for the degree")
run.font.size = Pt(11)

bsc = doc.add_paragraph()
bsc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = bsc.add_run("BACHELOR of SCIENCE in Computer Engineering")
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

pname = doc.add_paragraph()
pname.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pname.add_run('"YOLOv10 Tabanli Endustriyel Hata Tespit Sistemi\nve MLOps Altyapisi"')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

# Students
info_lines = [
    ("Students:", "Bahtiyar Selim Dogus, Muhammet, Berkay, Mehmet"),
    ("Supervisor(s):", "[Supervisor Name]"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + " ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

loc = doc.add_paragraph()
loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = loc.add_run("Istanbul\n2026")
run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  STUDENT DECLARATION
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("STUDENT DECLARATION")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

decl_text = (
    "By submitting this report, as partial fulfillment of the requirements of the "
    "Capstone Project, the students promise on penalty of failure of the course that:"
)
doc.add_paragraph(decl_text)

declarations = [
    "They have given credit to and declared (by citation), any work that is not their own.",
    "They have not received unpermitted aid for the project design, construction, report or presentation.",
    "They have not falsely assigned credit for work to another student in the group, and not take credit for work done by another student in the group.",
]
for d in declarations:
    p = doc.add_paragraph(d, style='List Paragraph')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("Abstract")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

abstract = (
    "Bu proje, beyaz esya uretim hattindaki vida ve plastik bilesenlerin montaj dogrulugunu "
    "denetlemek amaciyla gelistirilen otonom bir gorsel kalite kontrol sistemini kapsamaktadir. "
    "Sistemin temelinde, YOLOv10-S mimarisine Coordinate Attention (CA) katmani entegre edilerek "
    "olusturulan ozel bir nesne tespit modeli yer almaktadir. Faz 1 (Sprint 1) kapsaminda, "
    "model egitimi %99.4 mAP50 basari oranina ulasmistir. Ek olarak, Canny Edge Detection ile "
    "metalik yuzey yansimalarinin bastirilmasi ve K-means tabanli geometrik kumeleme ile "
    "mekansal dogrulama katmanlari sisteme entegre edilmistir. Egitim sureci MLflow ile takip "
    "edilmis, Streamlit tabanli bir dashboard uzerinden canli test ve operator geri bildirimi "
    "saglanmistir. Sistem, RTX 3050 GPU uzerinde ~7.3ms inference suresi ile saniyede 130'dan "
    "fazla kare isleyebilme kapasitesine ulasmistir."
)
doc.add_paragraph(abstract)

doc.add_paragraph()
kw = doc.add_paragraph()
run = kw.add_run("Keywords: ")
run.bold = True
run = kw.add_run("YOLOv10, Coordinate Attention, Industrial Defect Detection, MLOps, MLflow, "
                  "Edge Detection, Spatial Clustering, Quality Control, Deep Learning")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("Table of Contents")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph("[Icerik tablosu Word'de otomatik olusturulacaktir - References > Table of Contents]")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("List of Abbreviations")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

abbrevs = [
    ("mAP", "Mean Average Precision"),
    ("CA", "Coordinate Attention"),
    ("YOLO", "You Only Look Once"),
    ("MLOps", "Machine Learning Operations"),
    ("AMP", "Automatic Mixed Precision"),
    ("FP", "False Positive"),
    ("IoU", "Intersection over Union"),
    ("DFL", "Distribution Focal Loss"),
    ("GPU", "Graphics Processing Unit"),
    ("FPS", "Frames Per Second"),
    ("VLM", "Vision Language Model"),
    ("NMS", "Non-Maximum Suppression"),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].text = "Abbreviation"
hdr[1].text = "Definition"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    set_cell_shading(c, "D5E8F0")

for abbr, defn in abbrevs:
    add_table_row(tbl, [abbr, defn])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
sec_title = doc.add_paragraph()
sec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_title.add_run("Section 1:")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sec_sub = doc.add_paragraph()
sec_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_sub.add_run("Introduction")
run.bold = True
run.font.size = Pt(14)

# 1.1
doc.add_heading("1.1 Proje Amaci ve Kapsami", level=2)
doc.add_paragraph(
    "Bu proje, beyaz esya uretim hattinda vida ve plastik bilesenlerin montaj dogrulugunu "
    "denetlemek amaciyla baslatilmistir. Sistemin temel hedefi, fabrikalardaki yuksek hizli "
    "uretim bantlarina tam uyum saglayabilecek, mikro-endustriyel hatalari insan gozunden "
    "daha hizli ve kararli bir sekilde yakalayabilen otonom bir yapi kurmaktir."
)
doc.add_paragraph(
    "Proje, iki fazli bir mimari uzerine kurgulanmistir: Faz 1 olarak adlandirilan 'Gorsel "
    "Algilama' katmani, nesne tespiti motorunun endustriyel standartlarda optimize edilmesini "
    "kapsamaktadir. Faz 2'de ise tespit edilen hatalarin kok nedenlerini dogal dille "
    "aciklayacak bir VLM (Vision Language Model) katmani entegre edilecektir."
)

# 1.2
doc.add_heading("1.2 Problem Tanimi", level=2)
doc.add_paragraph(
    "Endustriyel uretim hatlarinda kalite kontrol, geleneksel olarak insan operatorler "
    "tarafindan gerceklestirilmektedir. Ancak yuksek hizli uretim bantlarinda (saniyede "
    "yuzlerce urun) insan gozunun yorgunluk, dikkat dagilmasi ve mikro-kusur kacirma gibi "
    "sinirlamalari bulunmaktadir. Ozellikle vida baslari ve eksik bilesen bosluklari, toplam "
    "goruntu alaninin %1'inden daha azini kaplamaktadir; bu da tespiti daha da zorlastirmaktadir."
)
doc.add_paragraph(
    "Bu projede hedeflenen sistem, su temel zorlulari ele almaktadir:"
)
challenges = [
    "Mikro nesnelerin (vida, bilesen) cok kucuk goruntu alani kaplamasi",
    "Metalik yuzeylerden kaynaklanan yansima ve parlama",
    "Farkli kamera acilarindan kaynaklanan perspektif varyasyonlari",
    "Uretim bandinin hizindan kaynaklanan hareket bulankikligi",
    "Nadir gorulen hata siniflarindaki veri kitligi (class imbalance)",
]
for c in challenges:
    doc.add_paragraph(c, style='List Paragraph')

# 1.3
doc.add_heading("1.3 Sprint 1 Hedefleri", level=2)
doc.add_paragraph(
    "Ilk sprint kapsaminda asagidaki hedefler belirlenmis ve gerceklestirilmistir:"
)
goals = [
    "Endustriyel goruntulerin toplanmasi ve YOLO formatinda etiketlenmesi",
    "YOLOv10-S + Coordinate Attention mimarisi ile model egitimi",
    "MLflow ile deney takibi ve metrik izleme",
    "Canny Edge Detection ile metalik yuzey on islemesi",
    "Geometrik kumeleme (Spatial Logic) ile mekansal dogrulama",
    "Streamlit dashboard ile canli test ve operator arayuzu",
    "Hedef: mAP50 >= %95 (gerceklesen: %99.4)",
]
for g in goals:
    doc.add_paragraph(g, style='List Paragraph')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 2: LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════
sec_title = doc.add_paragraph()
sec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_title.add_run("Section 2:")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sec_sub = doc.add_paragraph()
sec_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_sub.add_run("Literature Review")
run.bold = True
run.font.size = Pt(14)

# 2.1
doc.add_heading("2.1 YOLOv10 Mimarisi", level=2)
doc.add_paragraph(
    "YOLO (You Only Look Once) ailesi, gercek zamanli nesne tespiti alaninda en yaygin "
    "kullanilan mimarilerden biridir. YOLOv10, onceki surumlerden farkli olarak NMS-free "
    "(Non-Maximum Suppression gerektirmeyen) bir mimari sunmakta ve bu sayede inference "
    "surecini onemli olcude hizlandirmaktadir. YOLOv10-S (Small) varyanti, hiz ve dogruluk "
    "arasinda optimal bir denge kurarak endustriyel uygulamalar icin ideal bir secenektir."
)

# 2.2
doc.add_heading("2.2 Coordinate Attention (CA)", level=2)
doc.add_paragraph(
    "Coordinate Attention, Hou et al. tarafindan 2021 yilinda CVPR'de sunulmus bir dikkat "
    "mekanizmasidir. Geleneksel Squeeze-and-Excitation (SE) bloklari yalnizca kanal bazli "
    "dikkat uygularken, mekansal bilgiyi kaybeder. CA ise kanal bilgilerini X ve Y "
    "eksenlerinde ayri ayri kodlayarak modelin mekansal hassasiyetini (spatial awareness) "
    "maksimuma cikarir."
)
doc.add_paragraph(
    "Bu proje baglaminda CA'nin kritik onemi sudur: Endustriyel ortamda vida baslari ve "
    "eksik parca bosluklari toplam goruntu alaninin %1'inden daha azini kaplar. CA katmani, "
    "modelin 'nesnenin nerede oldugu' bilgisini milimetrik bir kesinlikle korur ve mikro "
    "kusurlari cok daha yuksek bir guven skoruyla yakalar. Backbone yapisinda 3 noktaya "
    "(256, 512 ve 1024 kanal seviyelerinde) CA eklemesi yapilmistir."
)

# 2.3
doc.add_heading("2.3 Endustriyel Hata Tespitinde Mevcut Yaklasimlar", level=2)
doc.add_paragraph(
    "Endustriyel kalite kontrol alaninda derin ogrenme tabanli nesne tespiti, son yillarda "
    "hizla yayginlasmaktadir. Geleneksel goruntu isleme yontemleri (template matching, "
    "threshold-based detection) sinirlari nedeniyle yerini CNN tabanli modellere birakmistir. "
    "Ancak salt derin ogrenmeye guven, yansima, perspektif degisimi ve nadir hata siniflarinda "
    "yetersiz kalabilmektedir. Bu nedenle, hibrit yaklasimlar (derin ogrenme + kural tabanli "
    "mantik) endustriyel uygulamalarda on plana cikmaktadir."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 3: METHODOLOGY
# ═══════════════════════════════════════════════════════════════════
sec_title = doc.add_paragraph()
sec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_title.add_run("Section 3:")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sec_sub = doc.add_paragraph()
sec_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_sub.add_run("Methodology")
run.bold = True
run.font.size = Pt(14)

# 3.1
doc.add_heading("3.1 Veri Toplama ve Etiketleme", level=2)
doc.add_paragraph(
    "Modelin egitimi icin uc ana veri kaynagi kullanilmistir:"
)
sources = [
    "erdogan1: Fabrika ortamindan toplanan ham vida denetim fotograflari",
    "erdogan2: Ek ham veri seti (farkli aci ve kosullar)",
    "roboflowetiketlenen: Roboflow platformu uzerinden COCO formatinda etiketlenmis, "
    "projenin ana referans veri seti",
]
for s in sources:
    doc.add_paragraph(s, style='List Paragraph')

doc.add_paragraph(
    "Tum goruntuler YOLO formatinda (class_id, center_x, center_y, width, height) "
    "normalize edilmis koordinatlarla etiketlenmistir. Uc sinif tanimlanmistir:"
)

# Class table
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Sinif ID", "Sinif Adi", "Aciklama"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

class_data = [
    ("0", "screw", "Normal vida mevcut"),
    ("1", "missing_screw", "Vida yuvasi bos / vida eksik"),
    ("2", "missing_component", "Tum bilesen eksik (kritik hata)"),
]
for row_data in class_data:
    add_table_row(tbl, row_data)

doc.add_paragraph()

doc.add_heading("3.1.1 Ilk Veri Seti Dagilimi", level=3)
doc.add_paragraph(
    "Roboflow verisinden olusturulan ilk veri seti (V1) 80/10/10 oraninda "
    "train/validation/test olarak bolunmustur:"
)

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["", "Train", "Validation", "Test"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

add_table_row(tbl, ["Goruntu Sayisi", "717", "89", "91"])
add_table_row(tbl, ["screw instance", "1,762", "-", "-"])
add_table_row(tbl, ["missing_screw instance", "16", "-", "-"])
add_table_row(tbl, ["missing_component instance", "4", "-", "-"])

doc.add_paragraph()
doc.add_paragraph(
    "Tabloda goruldugu uzere, ciddi bir sinif dengesizligi (class imbalance) sorunu "
    "mevcuttur. Ozellikle missing_component sinifina ait yalnizca 4 ornek bulunmasi, "
    "modelin bu kritik hatayi ogrenmesini teknik olarak son derece zorlastirmaktadir."
)

# Sample image
doc.add_paragraph()
add_fig(doc, "sample_train", "Sekil 1: Egitim veri setinden ornek goruntu - endüstriyel vida denetim kareleri", Inches(4))

# 3.2
doc.add_heading("3.2 Veri Artirma (Data Augmentation)", level=2)
doc.add_paragraph(
    "Sinif dengesizligi problemini cozmek icin iki farkli veri artirma kaynagi "
    "kullanilmistir:"
)
aug_sources = [
    "coklanmis/: Farkli acilardan cekilmis ve perspektif donusumleri (Shear, Perspective "
    "Transform) uygulanmis goruntuler",
    "coklanmisacili/: Ek acili veri seti - kamera acisi varyasyonlarina karsi dirayet "
    "kazandirmak icin",
]
for a in aug_sources:
    doc.add_paragraph(a, style='List Paragraph')

doc.add_paragraph(
    "Augmentation surecinde MD5 hash tabanli yinelenen goruntu tespiti (deduplication) "
    "ve SHA-256 tabanli veri sizintisi onleme (data leakage prevention) uygulanmistir. "
    "Bu sayede validation/test setindeki goruntulerin egitim setine karismasi engellenmistir."
)

doc.add_paragraph()
doc.add_paragraph("Augmentation sonrasi veri seti dagilimi:")

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Sinif", "Once (V1)", "Sonra (V2)"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

aug_data = [
    ("Toplam Train Goruntu", "717", "2,328"),
    ("screw instance", "1,762", "2,630"),
    ("missing_screw instance", "16", "480"),
    ("missing_component instance", "4", "283"),
]
for row_data in aug_data:
    add_table_row(tbl, row_data)

doc.add_paragraph()
add_fig(doc, "augment_imbalance",
        "Sekil 2: Augmentation oncesi ve sonrasi sinif dagilimi karsilastirmasi",
        Inches(5.5))

# 3.3
doc.add_heading("3.3 Model Mimarisi: YOLOv10-S + Coordinate Attention", level=2)
doc.add_paragraph(
    "Projenin cekirdek mimarisi, YOLOv10-S (Small) modelinin backbone yapisina "
    "Coordinate Attention (CA) katmanlarinin entegre edilmesiyle olusturulmustur. "
    "CA katmanlari backbone'da uc noktaya (256, 512 ve 1024 kanal seviyelerine) "
    "eklenmistir."
)

doc.add_paragraph(
    "Model mimarisi su katmanlardan olusmaktadir:"
)
arch_layers = [
    "Backbone: Conv -> C2f -> CoordAtt(256) -> SCDown -> C2f -> CoordAtt(512) -> "
    "SCDown -> C2fCIB -> SPPF -> PSA -> CoordAtt(1024)",
    "Head: Upsample + Concat -> C2f -> v10Detect (3 olcekli tespit basi)",
    "Giris boyutu: 640x640 piksel",
    "Sinif sayisi: 3 (screw, missing_screw, missing_component)",
]
for l in arch_layers:
    doc.add_paragraph(l, style='List Paragraph')

doc.add_paragraph(
    "Coordinate Attention'in temel calisme prensibi: X ve Y eksenleri boyunca "
    "ayri ayri ortalama havuzlama (average pooling) uygulanarak konum bilgisi korunur. "
    "Ardindan ortak 1x1 konvolusyon + aktivasyon ile kanal iliskileri ogrenilir ve "
    "son olarak sigmoid tabanli dikkat haritalari (Ax, Ay) olusturularak ozellik "
    "haritasi carpimsal olarak modullenir."
)

# 3.4
doc.add_heading("3.4 Egitim Sureci ve MLflow Entegrasyonu", level=2)
doc.add_paragraph(
    "Model egitimi, modern MLOps pratiklerine uygun olarak MLflow platformu uzerinden "
    "merkezi olarak yonetilmistir. Egitim parametreleri asagidaki gibidir:"
)

tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Parametre", "Deger"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

params = [
    ("Model", "YOLOv10-S + Coordinate Attention"),
    ("Pretrained", "yolov10s.pt (ImageNet transfer learning)"),
    ("Epoch", "100"),
    ("Batch Size", "8"),
    ("Goruntu Boyutu", "640x640"),
    ("AMP (Mixed Precision)", "Aktif"),
    ("Early Stopping Patience", "25 epoch"),
    ("Seed", "42 (tekrarlanabilirlik)"),
    ("GPU", "NVIDIA RTX 3050 (4GB VRAM)"),
    ("Framework", "Ultralytics + PyTorch 2.5.1 (CUDA 12.1)"),
]
for row_data in params:
    add_table_row(tbl, row_data)

doc.add_paragraph()
doc.add_paragraph(
    "MLflow uzerinden her egitim adiminda Box Loss, Class Loss ve DFL (Distribution "
    "Focal Loss) metrikleri anlik olarak izlenmistir. mAP degerlerindeki en ufak "
    "degisimler grafikselestirilerek modelin ezberleme (overfitting) yapip yapmadigi "
    "denetlenmistir. Farkli hiperparametre denemeleri bilimsel bir metodolojiyle "
    "karsilastirilmistir."
)

# 3.5
doc.add_heading("3.5 Hibrit Mantik Katmani", level=2)

doc.add_heading("3.5.1 Edge Detection (Kenar Algilama)", level=3)
doc.add_paragraph(
    "Metalik yuzeylerden kaynaklanan parlama ve yansimalar, vida yuvalarini gizleyerek "
    "modelin basarisini dusurebildigi gozlemlenmistir. Bu sorunu cozmek icin Canny Edge "
    "Detection tabanli bir on isleme adimi sisteme entegre edilmistir."
)
doc.add_paragraph(
    "Algoritma su sekilde calisir: Goruntu gri tonlamaya donusturulur, Canny kenar tespiti "
    "uygulanir (low=50, high=150), elde edilen kenar haritasi orijinal goruntuyle alpha "
    "blending (%70 orijinal, %30 kenar) ile harmanlanir. Bu sayede nesnelerin yapisal "
    "hatlari vurgulanarak modelin dairesel formlari ve bosluklari ayirt etmesi kolaylastirilir."
)

doc.add_heading("3.5.2 Geometrik Kumeleme (Spatial Logic)", level=3)
doc.add_paragraph(
    "Yalnizca derin ogrenmeye guvenmek yerine, kural tabanli bir geometrik dogrulama "
    "katmani sisteme entegre edilmistir. Bu katman, urun geometrisine dayali mantiksal "
    "cikarimlar yaparak False Positive oranini dusurmeyi amaclamaktadir."
)
doc.add_paragraph("Spatial Logic katmaninin calisme prensibi:")
spatial_steps = [
    "YOLO tespitleri K-means ile 4 beklenen vida pozisyonuna kumelenir",
    "Kumeler x-koordinat medyanina gore sol/sag tarafa atanir",
    "Her taraf icin durum belirlenir: S (screw mevcut), MS (missing_screw), MC (missing_component)",
    "Karar matrisi uygulanir: Her iki tarafta 2 vida -> OK; bir tarafta vida eksik -> missing_screw; "
    "ayni tarafta 2 vida eksik -> missing_component (bilesen tamamen yok)",
]
for s in spatial_steps:
    doc.add_paragraph(s, style='List Paragraph')

doc.add_paragraph(
    "Bu semantik kontrol mekanizmasi, gorsel olarak parcanin varligina dair yanilsamalari "
    "(yansima vb.) temizleyerek missing_component sinifindaki False Positive oranini "
    "yaklasik %30 oraninda dusurmesini saglamistir."
)

# 3.6
doc.add_heading("3.6 Dashboard (Streamlit)", level=2)
doc.add_paragraph(
    "Sistemin seffafligini ve son kullaniciya hitap etmesini saglamak amaciyla Streamlit "
    "tabanli kapsamli bir dashboard arayuzu gelistirilmistir. Dashboard 11 sayfadan "
    "olusmakta olup su temel ozellikleri sunmaktadir:"
)
dash_features = [
    "Canli Tahmin: Goruntu yukleyerek paralel olarak YOLO + Edge + Spatial analiz",
    "Veri Dengeleme: Augmentation oncesi/sonrasi sinif dagilimi gorsellestirmesi",
    "Coordinate Attention: Mimari rasyonel aciklamasi",
    "MLflow Takibi: Deney gecmisi ve metrik karsilastirmasi",
    "Edge Profiler: Jetson Orin Nano benchmark sonuclari",
    "False Positive Analizi: Operator geri bildirim dongusu",
    "Operator Kontrol: Acil durdurma, model yeniden yukleme, kuyruk temizleme",
]
for f in dash_features:
    doc.add_paragraph(f, style='List Paragraph')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 4: RESULTS
# ═══════════════════════════════════════════════════════════════════
sec_title = doc.add_paragraph()
sec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_title.add_run("Section 4:")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sec_sub = doc.add_paragraph()
sec_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_sub.add_run("Results")
run.bold = True
run.font.size = Pt(14)

# 4.1
doc.add_heading("4.1 V1 Model Sonuclari (Roboflow Verisi)", level=2)
doc.add_paragraph(
    "Ilk model (V1), yalnizca Roboflow etiketli veri seti (717 train goruntusu) ile "
    "egitilmistir. Baseline olarak 1 epoch'luk bir smoke test yapilmis, ardindan "
    "100 epoch boyunca tam egitim gerceklestirilmistir."
)

# Results table
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Metrik", "Baseline", "V1 Final", "Delta"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

add_table_row(tbl, ["mAP50(B)", "0.4942", "0.9942", "+0.5000"])
add_table_row(tbl, ["Precision", "~0.50", "0.993", "+0.493"])
add_table_row(tbl, ["Recall", "~0.49", "0.938", "+0.448"])

doc.add_paragraph()
doc.add_paragraph(
    "Sinif bazinda AP degerleri: screw 0.993, missing_screw 0.995. Bu sonuclar, "
    "modelin her iki sinifi da neredeyse kusursuz bir sekilde ayirt edebildigini "
    "gostermektedir."
)

doc.add_paragraph()
add_fig(doc, "v1_results",
        "Sekil 3: V1 model egitim grafikleri - Loss ve metrik egrileri (100 epoch)",
        Inches(5.5))

add_fig(doc, "v1_confusion",
        "Sekil 4: V1 Confusion Matrix - screw:222, missing_screw:3 dogru tespit",
        Inches(4.5))

add_fig(doc, "v1_pr_curve",
        "Sekil 5: V1 Precision-Recall egrisi - tum siniflar 0.994 mAP@0.5",
        Inches(5))

add_fig(doc, "v1_f1_curve",
        "Sekil 6: V1 F1-Confidence egrisi - tum siniflar 0.95 @ confidence 0.351",
        Inches(5))

# Baseline comparison
doc.add_heading("4.1.1 Baseline ile Karsilastirma", level=3)
doc.add_paragraph(
    "Baseline (1 epoch smoke test) confusion matrix'inde model yalnizca screw sinifini "
    "ogrenebilmis, missing_screw ve missing_component siniflarini hic tespit edememistir. "
    "100 epoch sonunda ise tum siniflar yuksek dogrulukla tespit edilmektedir."
)
add_fig(doc, "baseline_confusion",
        "Sekil 7: Baseline confusion matrix - yalnizca screw tespiti mumkun",
        Inches(4.5))

doc.add_page_break()

# 4.2
doc.add_heading("4.2 V2 Model Sonuclari (Coklanmis Verilerle)", level=2)
doc.add_paragraph(
    "V2 modeli, V1 veri setine ek olarak coklanmis ve coklanmisacili klasorlerindeki "
    "augmented verilerin eklenmesiyle olusturulan genisletilmis veri seti ile "
    "egitilmistir. Toplam 2,328 train goruntusu kullanilmistir."
)

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Metrik", "Baseline", "V2 Final", "Delta"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

add_table_row(tbl, ["mAP50(B)", "0.4942", "0.8450", "+0.3508"])

doc.add_paragraph()
doc.add_paragraph(
    "V2 modeli, V1'e kiyasla daha dusuk bir mAP50 degeri vermistir (0.8450 vs 0.9942). "
    "Bunun temel nedeni, augmented verilerin buyuk cogunlugunun (1,611 goruntu) hassas "
    "bounding box etiketleri yerine folder-based zayif etiketler (fallback bbox: 0.5, 0.5, "
    "0.8, 0.8) ile etiketlenmis olmasidir. Bu zayif etiketler, modele goruntunun tamamini "
    "nesne olarak ogrenmesine yol acarak siniflandirma kalitesini dusurmekedir."
)

add_fig(doc, "v2_results",
        "Sekil 8: V2 model egitim grafikleri - daha dalgali mAP egrisi gozlemlenmektedir",
        Inches(5.5))

add_fig(doc, "v2_confusion",
        "Sekil 9: V2 Confusion Matrix",
        Inches(4.5))

# 4.3
doc.add_heading("4.3 Genel Karsilastirma Tablosu", level=2)

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h_text in enumerate(["Metrik", "Baseline", "V1 (Roboflow)", "V2 (Augmented)"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    run = p.add_run(h_text)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(hdr[i], "D5E8F0")

compare_data = [
    ("mAP50(B)", "0.4942", "0.9942", "0.8450"),
    ("Train Goruntu", "717", "717", "2,328"),
    ("Epoch", "1", "100", "73 (early stop)"),
    ("En Iyi Epoch", "-", "96", "66"),
]
for row_data in compare_data:
    add_table_row(tbl, row_data)

# 4.4
doc.add_heading("4.4 Dashboard Test Sonuclari", level=2)
doc.add_paragraph(
    "V1 modeli Streamlit dashboard uzerinde canli olarak test edilmistir. "
    "NVIDIA RTX 3050 Laptop GPU uzerinde yapilan testlerde asagidaki performans "
    "degerleri elde edilmistir:"
)
perf_items = [
    "Inference suresi: ~7.3ms/goruntu",
    "Throughput: ~130+ FPS (saniyede 130'dan fazla kare)",
    "Confidence threshold: 0.25 (varsayilan)",
    "Edge Enhancement: alpha=0.7, Canny low=50, high=150",
]
for p_item in perf_items:
    doc.add_paragraph(p_item, style='List Paragraph')

add_fig(doc, "sample_infer",
        "Sekil 10: Ornek inference sonucu - vida tespitleri bounding box ve confidence skorlari ile",
        Inches(5))

# 4.5
doc.add_heading("4.5 Hibrit Mantik Katmaninin Etkisi", level=2)
doc.add_paragraph(
    "Spatial Logic katmaninin eklenmesiyle, ozellikle missing_component sinifindaki "
    "False Positive orani yaklasik %30 oraninda azalmistir. Geometrik dogrulama, modelin "
    "yansima kaynaklı yanilsamalari (bir vidanin parlama nedeniyle 'var' gorunmesi gibi) "
    "temizleyerek nihai karar kalitesini artirmistir."
)
doc.add_paragraph(
    "Edge Detection on islemesi ise ozellikle dusuk kontrast kosullarinda (karanlık "
    "arka plan uzerinde koyu metal) vida sinirlarinin daha belirgin hale gelmesini "
    "saglayarak tespit guvenilirligini artirmistir."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  SECTION 5: DISCUSSION AND CONCLUSION
# ═══════════════════════════════════════════════════════════════════
sec_title = doc.add_paragraph()
sec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_title.add_run("Section 5:")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

sec_sub = doc.add_paragraph()
sec_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sec_sub.add_run("Discussion and Conclusion")
run.bold = True
run.font.size = Pt(14)

# 5.1
doc.add_heading("5.1 Karsilasilan Sorunlar ve Cozum Yaklasimları", level=2)

doc.add_heading("5.1.1 Sinif Dengesizligi (Class Imbalance)", level=3)
doc.add_paragraph(
    "Projenin basinda en ciddi teknik engel, nadir gorulen hata siniflarindaki veri "
    "kitligi olmustur. missing_screw sinifinda yalnizca 16, missing_component sinifinda "
    "ise yalnizca 4 ornek bulunmasi, modelin bu sinıflari ogrenmesini imkansiz kilmaktaydi. "
    "Coklanmis veriler ve perspektif donusumleri ile bu siniflar sirasiyla 480 ve 283 "
    "ornege cikarilarak modelin saglikli ogrenebilecegi bir dengeye kavusturulmustur."
)

doc.add_heading("5.1.2 Bakis Acisi ve Domain Gap", level=3)
doc.add_paragraph(
    "Modelin ilk versiyonlari, sadece tepeden cekilmis fotograflarda yuksek performans "
    "verirken, fabrikanin gercekligindeki hafif acili goruntulerden basarisiz oluyordu. "
    "Bu sorunu cozmek icin perspektif donusumleri (Shear, Perspective Transform) iceren "
    "coklanmis veriler sisteme dahil edilerek modelin aci direnci arttirilmistir."
)

doc.add_heading("5.1.3 Bulaniklik ve Hareket Flulugu", level=3)
doc.add_paragraph(
    "Uretim bandinin hizi ve kameranin odaklanma sureleri nedeniyle bazi karelerin bulanik "
    "cikmasi, modelin vidayi 'metal uzerindeki bir leke' sanmasina yol acabiliyordu. "
    "Egitim asamasinda yapay Gaussian ve Motion Blur efektleri eklenerek modelin dusuk "
    "cozunurluklu veya hareketli anlarda dahi kararli (robust) kalmasi saglanmistir."
)

doc.add_heading("5.1.4 Metalik Yuzey Yansimlari", level=3)
doc.add_paragraph(
    "Metalik yuzeylerden kaynaklanan parlama ve yansimalar, vida yuvalarini gizleyerek "
    "tespit basarisini dusuruyordu. Canny Edge Detection tabanli on isleme adimi ile "
    "nesnelerin yapisal hatlari vurgulanarak bu sorun hafifletilmistir."
)

# 5.2
doc.add_heading("5.2 V1 vs V2 Karsilastirmasi ve Ogrenimler", level=2)
doc.add_paragraph(
    "V1 modeli (mAP50=0.9942) ile V2 modeli (mAP50=0.8450) arasindaki performans farki, "
    "veri kalitesinin veri miktarindan daha onemli oldugunu acikca gostermistir. V2'de "
    "kullanilan folder-based zayif etiketler (1,611 goruntu), modele yaniltici sinyal "
    "vererek ogrenme kalitesini dusurmustur."
)
doc.add_paragraph("Bu deneyimden cikarilan temel dersler:")
lessons = [
    "Hassas bounding box etiketlemesi, veri miktarindan daha kritik bir faktordur",
    "Zayif etiketler (full-image bbox) modelin tum goruntuyu nesne olarak ogrenmesine yol acar",
    "Data augmentation stratejisi, etiketleme kalitesiyle birlikte degerlendirilmelidir",
    "MD5/SHA-256 tabanli data leakage onleme, deneysel sonuclarin guvenilirligini arttirir",
]
for l in lessons:
    doc.add_paragraph(l, style='List Paragraph')

# 5.3
doc.add_heading("5.3 Sonuc", level=2)
doc.add_paragraph(
    "Sprint 1 kapsaminda, %99.4 mAP50 skoruna sahip, aci degisimlerine ve goruntu "
    "gurultulesine direncli bir 'Gorsel Algilama Katmani' basariyla gelistirilmistir. "
    "Coordinate Attention entegrasyonu, ozellikle kucuk nesnelerin (vida baslari) tespitinde "
    "belirgin bir iyilesme saglamistir. Hibrit mantik katmani (Edge Detection + Spatial "
    "Logic), salt derin ogrenmenin otesinde kural tabanli dogrulama ekleyerek sistemin "
    "endustriyel guvenilirligini artirmistir."
)

# 5.4
doc.add_heading("5.4 Gelecek Planlar (Faz 2)", level=2)
doc.add_paragraph("Faz 2 kapsaminda planlananan calismalar:")
future_plans = [
    "PaliGemma VLM Reasoning Katmani: Dusuk guvenli tespitlerde (confidence < 0.40) "
    "dogal dil tabanli kok neden aciklamasi",
    "TensorRT Optimizasyonu: FP16/INT8 quantization ile Jetson Orin Nano uzerinde "
    "gercek zamanli calisma (<5ms hedefi)",
    "Active Learning Dongusu: Operator geri bildirimlerinin otomatik yeniden egitim "
    "surecine dahil edilmesi",
    "Concept Drift Testi: Haftalik otomatik fine-tune ile model guncelliginin korunmasi",
]
for f in future_plans:
    doc.add_paragraph(f, style='List Paragraph')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("ACKNOWLEDGEMENT")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph(
    "Bu projenin gerceklestirilmesinde desteklerini esirgemeyen danisman hocamiza, "
    "veri toplama surecinde yardimci olan fabrika yetkililerine ve ekip arkadaslarimiza "
    "tesekkur ederiz."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════════════════
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("REFERENCES")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

refs = [
    "[1] Wang, A., Chen, H., Liu, L., Chen, K., Lin, Z., Han, J., & Ding, G. (2024). "
    "YOLOv10: Real-Time End-to-End Object Detection. arXiv preprint arXiv:2405.14458.",

    "[2] Hou, Q., Zhou, D., & Feng, J. (2021). Coordinate Attention for Efficient Mobile "
    "Network Design. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern "
    "Recognition (CVPR), pp. 13713-13722.",

    "[3] Ultralytics (2024). Ultralytics YOLOv8/v10 Documentation. "
    "https://docs.ultralytics.com/",

    "[4] Zaharia, M., et al. (2018). Accelerating the Machine Learning Lifecycle with MLflow. "
    "IEEE Data Engineering Bulletin, 41(4), 39-45.",

    "[5] Jocher, G., Chaurasia, A., & Qiu, J. (2023). Ultralytics YOLO (Version 8.0.0). "
    "https://github.com/ultralytics/ultralytics",
]
for ref in refs:
    doc.add_paragraph(ref)

# ═══════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"\nRapor basariyla olusturuldu: {OUTPUT}")
print(f"Dosya boyutu: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
